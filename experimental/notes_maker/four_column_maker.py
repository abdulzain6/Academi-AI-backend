from docx import Document
from docxtpl import DocxTemplate
from pydantic import BaseModel
from typing import List
from .base import NotesMaker
import io
import os

class QuestionKeywordItem(BaseModel):
    question: str
    keyword: str

class DataModel(BaseModel):
    notes1: List[str]
    notes2: List[str]
    questions_keywords: List[QuestionKeywordItem]

class ContextModel(BaseModel):
    SUMMARY: str
    NAME: str
    TOPIC: str
    DATE: str
    
    
class QuestionsKeywordsNotesMaker(NotesMaker):

    def __init__(self, template_path=None):
        if template_path is None:
            script_dir = os.path.dirname(os.path.abspath(__file__))
            template_path = os.path.join(script_dir, "templates", "4_COLUMN.docx")
        self.template_path = template_path
        self.docxtpl = DocxTemplate(template_path)

    def make_notes(self, data: DataModel, context: ContextModel):
        docx = Document(self.template_path)
        self.populate_table1(docx, data.notes1, data.notes2)
        self.populate_table2(docx, data.questions_keywords)
        return self.render_and_return(docx, context)

    def populate_table1(self, docx, notes1, notes2):
        table1 = docx.tables[0]
        existing_rows_table1 = len(table1.rows) - 1

        for i in range(max(len(notes1), len(notes2))):  # Iterate for the maximum of the two list lengths
            if i < existing_rows_table1:
                row_cells = table1.rows[i + 1].cells
            else:
                row_cells = table1.add_row().cells

            if i < len(notes1):
                run = row_cells[0].paragraphs[0].runs[0] if row_cells[0].paragraphs else row_cells[0].paragraphs[0].add_run()
                run.text = notes1[i]
                run.bold = True

            if i < len(notes2):
                run = row_cells[1].paragraphs[0].runs[0] if row_cells[1].paragraphs else row_cells[1].paragraphs[0].add_run()
                run.text = notes2[i]
                run.bold = True

    def populate_table2(self, docx, questions_keywords):
        table2 = docx.tables[1]
        existing_rows_table2 = len(table2.rows) - 1

        for i, item in enumerate(questions_keywords):
            if i < existing_rows_table2:
                row_cells = table2.rows[i + 1].cells
            else:
                row_cells = table2.add_row().cells

            for j, key in enumerate(['question', 'keyword']):
                run = row_cells[j].paragraphs[0].runs[0] if row_cells[j].paragraphs else row_cells[j].paragraphs[0].add_run()
                run.text = getattr(item, key)
                run.bold = True

    def render_and_return(self, docx, context: ContextModel):
        docx_stream = self.change_to_file_obj(docx)
        docxtpl = DocxTemplate(docx_stream)
        docxtpl.render(context.model_dump())  # Render other placeholders
        return self.change_to_file_obj(docxtpl)

    def change_to_file_obj(self, arg0):
        result = io.BytesIO()
        arg0.save(result)
        result.seek(0)
        return result

if __name__ == '__main__':
    # Define the input data and context
    data_input = {
        'notes1': ['Note 1A', 'Note 2A', 'Note 3A'],
        'notes2': ['Note 1B', 'Note 2B', 'Note 3B'],
        'questions_keywords': [
            {'question': 'Question 1', 'keyword': 'Keyword 1'},
            {'question': 'Question 2', 'keyword': 'Keyword 2'},
            {'question': 'Question 3', 'keyword': 'Keyword 3'},
            # ... add more data as needed
        ]
    }
    context_input = {
        "SUMMARY": "SUMMARY",
        "NAME": "NAME",
        "TOPIC": "TOPIC",
        "DATE": "DATE"
    }

    # Validate and convert input data and context using Pydantic models
    data_model = DataModel(**data_input)
    context_model = ContextModel(**context_input)

    # Create an instance of QuestionsKeywordsNotesMaker
    notes_maker = QuestionsKeywordsNotesMaker()

    # Call the make_notes method to generate the notes
    file_object = notes_maker.make_notes(data=data_model, context=context_model)

    # Save the resulting file object to a file
    with open("output.docx", "wb") as file:
        file.write(file_object.read())
