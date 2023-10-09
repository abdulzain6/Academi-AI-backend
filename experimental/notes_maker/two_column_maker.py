from typing import List
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt
from pydantic import BaseModel, Field
import io
import os
from .base import NotesMaker


class ContextModelTwoColumn(BaseModel):
    TITLE: str
    TOPIC: str
    SUMMARY: str
    
class QuestionDetailItem(BaseModel):
    question: str
    details: str

class DataModelTwoColumn(BaseModel):
    items: List[QuestionDetailItem] = Field(..., alias='data')
    


class QuestionsDetailsSummaryNotesMaker(NotesMaker):
    def __init__(self, template_path = None):
        if template_path is None:
            script_dir = os.path.dirname(os.path.abspath(__file__))
            template_path = os.path.join(script_dir, "templates", "2_COLUMN.docx")
        
        self.template_path = template_path
        self.docxtpl = DocxTemplate(template_path)
        
    def make_notes(self, data: DataModelTwoColumn, context: ContextModelTwoColumn):
        docx = Document(self.template_path)
        table = docx.tables[0]
        existing_rows = len(table.rows) - 3

        for i, row_data in enumerate(data.items):
            if i < existing_rows:
                row_cells = table.rows[i + 3].cells
            else:
                row_cells = table.add_row().cells

            for j, attr_name in enumerate(['question', 'details']):
                paragraph = row_cells[j].paragraphs[0] if row_cells[j].paragraphs else row_cells[j].add_paragraph()
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.text = getattr(row_data, attr_name)
                run.bold = True
                run.font.size = Pt(10)

        docx_stream = self.change_to_bytes_obj(docx)
        docxtpl = DocxTemplate(docx_stream)
        docxtpl.render(context=context.model_dump())  
        return self.change_to_bytes_obj(docxtpl)

    def change_to_bytes_obj(self, arg0):
        result = io.BytesIO()
        arg0.save(result)
        result.seek(0)
        return result

if __name__ == '__main__':
    # Define the data
    data_items = [
        QuestionDetailItem(question="What is the capital of France?", details="Paris"),
        QuestionDetailItem(question="What is the capital of Germany?", details="Berlin"),
        # ... add more items as needed
    ]
    data = DataModelTwoColumn(data=data_items)
    
    # Define the context
    context = ContextModelTwoColumn(
        TITLE="Geography Questions",
        TOPIC="Capital Cities",
        SUMMARY="A list of questions and answers about capital cities."
    )
    
    # Create an instance of QuestionsDetailsSummaryNotesMaker
    notes_maker = QuestionsDetailsSummaryNotesMaker()
    
    # Generate the notes
    file_object = notes_maker.make_notes(data, context)
    
    # Optionally, save the generated notes to a file
    output_path = 'output.docx'
    with open(output_path, 'wb') as f:
        f.write(file_object.read())
    print(f'Document saved at: {output_path}')
