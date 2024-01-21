from typing import List
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt
from pydantic import BaseModel, Field
from .base import NotesMaker
from langchain.chains import LLMChain
from langchain.chat_models.base import BaseChatModel
from langchain.prompts import (
    ChatPromptTemplate,
    SystemMessagePromptTemplate,
    HumanMessagePromptTemplate,
)
from langchain.output_parsers import PydanticOutputParser, OutputFixingParser
import io
import os


class ContextModelTwoColumn(BaseModel):
    TITLE: str = Field(
        json_schema_extra={
            "description": "Title for the notes."
        }
    )
    TOPIC: str = Field(
        json_schema_extra={
            "description": "Topic for the notes."
        }
    )
    SUMMARY: str = Field(
        json_schema_extra={
            "description": "Summary for the notes"
        }
    )


class QuestionDetailItem(BaseModel):
    question: str = Field(
        json_schema_extra={
            "description": "Important Question / Main Idea / Vocabulary worth noting from the data."
        }
    )
    details: str = Field(
        json_schema_extra={
            "description": "Notes / Answer / Definition / Example / Sentence, about the chosen Question / Main Idea / Vocabulary"
        }
    )

class DataModelTwoColumn(BaseModel):
    items: List[QuestionDetailItem] = Field(..., alias="data", json_schema_extra={
        "description": "List of question answer pairs"
    })


class InputData(BaseModel):
    notes_metadata: ContextModelTwoColumn = Field(
        json_schema_extra={"description": "Extra information to put into the notes"}
    )
    notes: DataModelTwoColumn = Field(..., json_schema_extra={
        "description": "List of question answer pairs"
    })


class QuestionsDetailsSummaryNotesMaker(NotesMaker):
    def __init__(self, llm: BaseChatModel, template_path=None):
        if template_path is None:
            script_dir = os.path.dirname(os.path.abspath(__file__))
            template_path = os.path.join(script_dir, "templates", "2_COLUMN.docx")

        self.template_path = template_path
        self.docxtpl = DocxTemplate(template_path)
        self.llm = llm
        
    @staticmethod
    def get_schema():
        return InputData.model_json_schema()

    def make_notes_from_string(self, string: str, instructions: str) -> io.BytesIO:
        parser = PydanticOutputParser(pydantic_object=InputData)
        parser = OutputFixingParser.from_llm(parser=parser, llm=self.llm)
        prompt = ChatPromptTemplate(
            messages=[
                SystemMessagePromptTemplate.from_template(
                    """You are an AI designed to make cornell notes from text.
You will return json in the schema told to you (Important).
Failure to do this will cause an error.
You must also follow the instructions given to you
Only answer from the notes given to you. No making things up
You must pick every minute detail.
You must make super detailed and lenghty notes
"""
                ),
                HumanMessagePromptTemplate.from_template(
                    """
Use the following data to make the notes:
{data}
============

Follow the instructions below also:
============
{instructions}
============

The notes must be formatted in the schema below:
{format_instructions}

The notes in proper format (Failure causes big error):"""
                ),
            ],
            input_variables=[
                "data",
                "instructions"
            ],
            partial_variables={"format_instructions" : parser.get_format_instructions()}
        )
        chain = LLMChain(prompt=prompt, output_parser=parser, llm=self.llm)
        notes: InputData = chain.run(data=string, instructions=instructions)
        return self.make_notes(notes.notes, notes.notes_metadata)

    def make_notes(
        self, data: DataModelTwoColumn, context: ContextModelTwoColumn
    ) -> io.BytesIO:
        docx = Document(self.template_path)
        table = docx.tables[0]
        existing_rows = len(table.rows) - 3

        for i, row_data in enumerate(data.items):
            if i < existing_rows:
                row_cells = table.rows[i + 3].cells
            else:
                row_cells = table.add_row().cells

            for j, attr_name in enumerate(["question", "details"]):
                paragraph = (
                    row_cells[j].paragraphs[0]
                    if row_cells[j].paragraphs
                    else row_cells[j].add_paragraph()
                )
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.text = getattr(row_data, attr_name)
                run.bold = True
                run.font.size = Pt(10)

        docx_stream = self.change_to_bytes_obj(docx)
        docxtpl = DocxTemplate(docx_stream)
        docxtpl.render(context=context.model_dump())
        return self.change_to_bytes_obj(docxtpl)

    def change_to_bytes_obj(self, arg0: DocxTemplate) -> io.BytesIO:
        result = io.BytesIO()
        arg0.save(result)
        result.seek(0)
        return result

    def make_notes_from_dict(self, data_dict: str) -> io.BytesIO:
        input_data = InputData.model_validate(data_dict)
        return self.make_notes(input_data.notes, context=input_data.notes_metadata)



if __name__ == "__main__":
    from langchain.chat_models import ChatOpenAI
    
    note_maker = QuestionsDetailsSummaryNotesMaker(
        ChatOpenAI(temperature=0, openai_api_key="")
    )
    notes = note_maker.make_notes_from_string("""
World War II: An Epoch of Global Conflict and Change

The history of the 20th century is profoundly marked by the cataclysmic events of World War II, a global conflict that spanned six years, from 1939 to 1945. This war reshaped the world's geopolitical boundaries, altered global power dynamics, and set the stage for both the Cold War and the modern world order.

The war was principally a confrontation between two sets of powers: the Axis and the Allies. The Axis powers, chiefly comprised of Germany, Italy, and Japan, were driven by expansionist ambitions and authoritarian ideologies. Germany, under the rule of Adolf Hitler, played a pivotal role in the war's outbreak and progression. Hitler's aggressive policies and relentless pursuit of territorial expansion pushed Europe into a state of turmoil.

Japan's role in the Pacific theater was equally consequential. The Japanese attack on Pearl Harbor on December 7, 1941, marked a significant escalation of the war, drawing the United States into active combat. This surprise attack, aimed at neutralizing the American Pacific Fleet, stemmed from Japan's desire to dominate Southeast Asia without American intervention. The strategic motivations behind Japan's military actions in the Pacific, driven by economic and resource scarcities, illustrate the complex interplay of national interests that characterized the war.

The Allies, consisting of major powers like the United Kingdom, the United States, the Soviet Union, and France, among others, united to counter the aggression of the Axis powers. Their collaboration was marked by a shared commitment to defeating the forces of fascism and militarism.

Several key battles defined the trajectory of World War II. The Battle of Stalingrad was a turning point on the Eastern Front. The Soviet victory in this grueling battle marked the beginning of the decline for Nazi Germany. The Battle of Midway, a crucial naval conflict, signified a decisive shift in the Pacific, as American forces dealt a critical blow to the Japanese navy.

Another significant event was the Normandy Invasion, commonly known as D-Day, on June 6, 1944. This massive military operation, involving the landing of Allied forces on the beaches of Normandy, France, was a pivotal moment in liberating Western Europe from Nazi occupation.

The war's impact extended far beyond the battlefield. It brought about profound changes in political, social, and cultural realms. The Holocaust, a horrific genocide perpetrated by Nazi Germany, led to the systematic extermination of six million Jews, along with millions of others deemed undesirable by the Nazi regime. This atrocity highlighted the depths of human cruelty and the necessity for a robust international framework to prevent such crimes against humanity.

In the wake of the war, the United Nations was established, symbolizing the global community's aspiration for peace and cooperation. However, the world also witnessed the emergence of the Cold War, a period of geopolitical tension between the Soviet Union and the United States that lasted for decades.

As we delve deeper into the multifaceted narrative of World War II, it becomes evident that this conflict was not just a series of military engagements but also a clash of ideologies and civilizations. It challenged the notions of power, sovereignty, and human rights, leaving an indelible imprint on the course of human history.
""", "Dont miss anything")
    with open("file.docx", "wb") as fp:
        notes.seek(0)
        fp.write(notes.read())
