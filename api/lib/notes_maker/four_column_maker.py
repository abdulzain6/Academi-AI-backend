from docx import Document
from docxtpl import DocxTemplate
from pydantic import BaseModel, Field
from typing import List
from .base import NotesMaker
from langchain.chat_models.base import BaseChatModel
from langchain.prompts import (
    ChatPromptTemplate,
    SystemMessagePromptTemplate,
    HumanMessagePromptTemplate,
)
from langchain.chains import LLMChain
from langchain.output_parsers import PydanticOutputParser, OutputFixingParser

from datetime import date
import os, io

class QuestionKeywordItem(BaseModel):
    question: str  = Field(
        json_schema_extra={
            "description": "Important Question / Main Idea / Vocabulary worth noting from the data."
        }
    )
    answer: str = Field(
        json_schema_extra={
            "description": "Answers to the question"
        }
    )

class DataModel(BaseModel):
    notes1: List[str] = Field(
        json_schema_extra={
            "description": "List Number one of notes from the data."
        }
    )
    notes2: List[str] = Field(
        json_schema_extra={
            "description": "List Number two of notes from the data. Must be different from notes one."
        }
    )
    questions_keywords: List[QuestionKeywordItem] = Field(
        json_schema_extra={
            "description": "Question/Keyword pairs from the data."
        }
    )

class ContextModel(BaseModel):
    NAME: str = Field(
        "John Doe",
        json_schema_extra={
            "description": "Name of the student"
        }
    )
    DATE: str = Field(
        f"{date.today()}",
        json_schema_extra={
            "description": "Date of note"
        }
    )
    TOPIC: str = Field(
        json_schema_extra={
            "description": "Topic for the notes."
        }
    )
    SUMMARY: str = Field(
        json_schema_extra={
            "description": "Summary for the notes"
        }
    )
    
class InputData(BaseModel):
    notes_metadata: ContextModel = Field(
        json_schema_extra={"description": "Extra information to put into the notes"}
    )
    notes: DataModel = Field(..., json_schema_extra={
        "description": "List of question answer pairs"
    })
    
class QuestionsKeywordsNotesMaker(NotesMaker):
    def __init__(self, llm: BaseChatModel, template_path=None):
        if template_path is None:
            script_dir = os.path.dirname(os.path.abspath(__file__))
            template_path = os.path.join(script_dir, "templates", "4_COLUMN.docx")
        
        self.llm = llm
        self.template_path = template_path
        self.docxtpl = DocxTemplate(template_path)
        
    @staticmethod
    def get_schema():
        return InputData.model_json_schema()
    
    def make_notes_from_dict(self, data_dict: str) -> io.BytesIO:
        input_data = InputData.model_validate(data_dict)
        return self.make_notes(input_data.notes, context=input_data.notes_metadata)

    def make_notes(self, data: DataModel, context: ContextModel):
        docx = Document(self.template_path)
        self.populate_table1(docx, data.notes1, data.notes2)
        self.populate_table2(docx, data.questions_keywords)
        return self.render_and_return(docx, context)

    def populate_table1(self, docx: Document, notes1: List[str], notes2: List[str]) -> None:
        table1 = docx.tables[0]
        existing_rows_table1 = len(table1.rows) - 1

        for i in range(max(len(notes1), len(notes2))):  # Iterate for the maximum of the two list lengths
            row_cells = table1.rows[i + 1].cells if i < existing_rows_table1 else table1.add_row().cells

            for col_idx, notes in enumerate([notes1, notes2]):
                if i < len(notes):
                    if row_cells[col_idx].paragraphs:
                        run = row_cells[col_idx].paragraphs[0].runs[0] if row_cells[col_idx].paragraphs[0].runs else row_cells[col_idx].paragraphs[0].add_run()
                    else:
                        run = row_cells[col_idx].add_paragraph().add_run()

                    run.text = notes[i]
                    run.bold = True

    def populate_table2(self, docx: Document, questions_keywords: List[QuestionKeywordItem]) -> None:
        table2 = docx.tables[1]
        existing_rows_table2 = len(table2.rows) - 1

        for i, item in enumerate(questions_keywords):
            row_cells = table2.rows[i + 1].cells if i < existing_rows_table2 else table2.add_row().cells

            for j, key in enumerate(['question', 'answer']):
                if row_cells[j].paragraphs:
                    run = row_cells[j].paragraphs[0].runs[0] if row_cells[j].paragraphs[0].runs else row_cells[j].paragraphs[0].add_run()
                else:
                    run = row_cells[j].add_paragraph().add_run()

                run.text = getattr(item, key)
                run.bold = True

    def render_and_return(self, docx, context: ContextModel):
        docx_stream = self.change_to_file_obj(docx)
        docxtpl = DocxTemplate(docx_stream)
        docxtpl.render(context.model_dump())  # Render other placeholders
        return self.change_to_file_obj(docxtpl)

    def change_to_file_obj(self, arg0):
        result = io.BytesIO()
        arg0.save(result)
        result.seek(0)
        return result

    def make_notes_from_string(self, string: str, instructions: str) -> io.BytesIO:
        parser = PydanticOutputParser(pydantic_object=InputData)
        parser = OutputFixingParser.from_llm(parser=parser, llm=self.llm)
        prompt = ChatPromptTemplate(
            messages=[
                SystemMessagePromptTemplate.from_template(
                    """You are an AI designed to make cornell notes from text.
You will return json in the schema told to you (Important).
Failure to do this will cause an error.
You must also follow the instructions given to you
Only answer from the notes given to you. No making things up
You must pick every minute detail.
You must make super detailed and lenghty notes
"""
                ),
                HumanMessagePromptTemplate.from_template(
                    """
Use the following data to make the notes:
{data}
============

Follow the instructions below also:
============
{instructions}
============

The notes must be formatted in the schema below:
{format_instructions}

The notes in proper format (Failure causes big error):"""
                ),
            ],
            input_variables=[
                "data",
                "instructions"
            ],
            partial_variables={"format_instructions" : parser.get_format_instructions()}
        )
        chain = LLMChain(prompt=prompt, output_parser=parser, llm=self.llm)
        notes: InputData = chain.run(data=string, instructions=instructions)
        return self.make_notes(notes.notes, notes.notes_metadata)
